# 2021 Copyright Dinu Ion-Irinel
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

# initialize of the document
document = Document()

# intro informations 
name = input("Hello, what is your name? : ")
phone_number = input("What's your phone_number: ")
email = input("Enter your email: ")
image_tag = input("Enter the name of image: ")


document.add_picture('./' + image_tag, width=Inches(2.0))
document.add_heading(name + '\n')
document.add_paragraph(phone_number + ' | ' + email)


# about informations
document.add_heading("About Me")
informations = input("Tell us something about you...")
document.add_paragraph(informations)

# experience area
document.add_heading("Experience:")
while True:
    first_user_response = input("Do you have more experience? ")
    if first_user_response.lower() == "no":
        break
    else:
        company = input("Enter a company: ")
        perioad = input("Enter the period: ")
        describe_experience = input("Describe your experience: ")
        experiences_paragraph = document.add_paragraph()
        experiences_paragraph.add_run(company + " ").bold = True
        experiences_paragraph.add_run(perioad + "\n").italic = True
        experiences_paragraph.add_run(describe_experience)

# skills area
document.add_heading("Skills")
while True:
    second_user_response = input("Do you have more skills? ")
    if second_user_response.lower() == "no":
        break
    else:
        skill = input("Enter your skill: ")
        skills_paragraph = document.add_paragraph()
        skills_paragraph.style = 'List Bullet'
        skills_paragraph.add_run(skill)


# languages area
document.add_heading("Language")
while True:
    third_user_response = input("Do you know more language? ")
    if third_user_response.lower() == 'no':
        break
    else:
        language = input("Enter a new language: ")
        languages_paragraph = document.add_paragraph()
        languages_paragraph.style = 'List Bullet'
        languages_paragraph.add_run(language)

# save the document
document.save("resume.docx")

